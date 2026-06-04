from fastapi import APIRouter, Depends, HTTPException, status, Request, Query
from fastapi.responses import StreamingResponse
from datetime import datetime
import io
import docx
from docx.shared import Pt, RGBColor, Inches
from fpdf import FPDF
import logging
from typing import Optional

from app.middleware.auth_middleware import get_current_user
from app.database import get_collection
from app.utils.jwt_handler import verify_token

router = APIRouter(prefix="/api/export", tags=["Report Exports"])
logger = logging.getLogger(__name__)

async def get_user_for_export(request: Request, token: Optional[str] = Query(None)):
    tok = token
    if not tok:
        auth_header = request.headers.get("Authorization")
        if auth_header and auth_header.startswith("Bearer "):
            tok = auth_header.split(" ")[1]
            
    if not tok:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Authentication token is missing."
        )
        
    payload = verify_token(tok)
    if not payload:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Session expired or invalid token."
        )
        
    email = payload.get("sub")
    if not email:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid token subject."
        )
        
    users_collection = get_collection("users")
    user = await users_collection.find_one({"email": email})
    if not user:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="User not found."
        )
        
    user["_id"] = str(user["_id"])
    return user

class LegalReportPDF(FPDF):
    def header(self):
        # Header banner
        self.set_font('Helvetica', 'B', 8)
        self.set_text_color(100, 110, 130)
        self.cell(0, 5, 'VERDICTIQ AI LEGAL OPERATING SYSTEM // CASE INTELLIGENCE REPORT', border=0, ln=1, align='L')
        self.line(10, 15, 200, 15)
        self.ln(5)

    def footer(self):
        # Position at 1.5 cm from bottom
        self.set_y(-15)
        self.line(10, 280, 200, 280)
        self.set_font('Helvetica', 'I', 8)
        self.set_text_color(120, 120, 120)
        # Page number
        self.cell(0, 10, f'Page {self.page_no()}/{{nb}} // SECURE QUANTUM ENCRYPTED REPORT', align='C')

@router.get("/pdf/{workspace_id}")
async def export_pdf(workspace_id: str, current_user: dict = Depends(get_user_for_export)):
    # 1. Fetch metadata and report from database
    workspace = await get_collection("workspaces").find_one({"workspace_id": workspace_id})
    report = await get_collection("agent3_final_reports").find_one({"workspace_id": workspace_id})

    if not workspace or not report:
        raise HTTPException(status_code=404, detail="Workspace metadata or final report not found. Run analysis first.")

    # 2. Build PDF Document
    pdf = LegalReportPDF()
    pdf.alias_nb_pages()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=20)

    # Document Title Block
    pdf.set_font('Helvetica', 'B', 20)
    pdf.set_text_color(30, 41, 59) # Slate 800
    pdf.cell(0, 10, "LEGAL INTELLIGENCE AUDIT", ln=1, align='C')
    
    pdf.set_font('Helvetica', 'B', 12)
    pdf.set_text_color(79, 140, 255) # Brand Blue
    pdf.cell(0, 8, f"CASE: {workspace.get('case_title', 'Unnamed Case').upper()}", ln=1, align='C')
    pdf.ln(5)

    # Metadata Table Box
    pdf.set_fill_color(240, 245, 255)
    pdf.rect(10, 40, 190, 40, 'F')
    
    pdf.set_xy(12, 42)
    pdf.set_font('Helvetica', 'B', 9)
    pdf.set_text_color(60, 60, 60)
    pdf.cell(40, 6, "Case Type:", 0, 0)
    pdf.set_font('Helvetica', '', 9)
    pdf.cell(50, 6, str(workspace.get('case_type', 'N/A')), 0, 0)
    
    pdf.set_font('Helvetica', 'B', 9)
    pdf.cell(40, 6, "Counsel Side:", 0, 0)
    pdf.set_font('Helvetica', '', 9)
    pdf.cell(50, 6, str(workspace.get('lawyer_side', 'N/A')), 0, 1)

    pdf.set_x(12)
    pdf.set_font('Helvetica', 'B', 9)
    pdf.cell(40, 6, "Client:", 0, 0)
    pdf.set_font('Helvetica', '', 9)
    pdf.cell(50, 6, str(workspace.get('client_name', 'N/A')), 0, 0)
    
    pdf.set_font('Helvetica', 'B', 9)
    pdf.cell(40, 6, "Opposing Party:", 0, 0)
    pdf.set_font('Helvetica', '', 9)
    pdf.cell(50, 6, str(workspace.get('opposing_party', 'N/A')), 0, 1)

    pdf.set_x(12)
    pdf.set_font('Helvetica', 'B', 9)
    pdf.cell(40, 6, "Report Date:", 0, 0)
    pdf.set_font('Helvetica', '', 9)
    pdf.cell(50, 6, datetime.now().strftime("%Y-%m-%d %H:%M"), 0, 0)
    
    pdf.set_font('Helvetica', 'B', 9)
    pdf.cell(40, 6, "Overall Case Strength:", 0, 0)
    pdf.set_font('Helvetica', 'B', 9)
    pdf.set_text_color(220, 50, 50) if "high" in report.get('case_strength', '').lower() else pdf.set_text_color(34, 197, 94)
    pdf.cell(50, 6, str(report.get('case_strength', 'N/A')), 0, 1)
    
    pdf.set_xy(10, 85)
    pdf.ln(5)

    # Executive Summary Section
    pdf.set_font('Helvetica', 'B', 12)
    pdf.set_text_color(30, 41, 59)
    pdf.cell(0, 8, "1. EXECUTIVE SUMMARY", ln=1)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(2)
    
    pdf.set_font('Helvetica', '', 10)
    pdf.set_text_color(50, 50, 50)
    pdf.multi_cell(0, 5.5, str(report.get('executive_summary', 'No summary available.')))
    pdf.ln(6)

    # Evidence Auditing Sections
    pdf.set_font('Helvetica', 'B', 12)
    pdf.set_text_color(30, 41, 59)
    pdf.cell(0, 8, "2. EVIDENTIARY AUDIT", ln=1)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(2)
    
    pdf.set_font('Helvetica', 'B', 10)
    pdf.cell(0, 6, "Strong Evidence:", ln=1)
    pdf.set_font('Helvetica', '', 9.5)
    strong_ev = report.get('strongest_evidence', [])
    if strong_ev:
        for idx, item in enumerate(strong_ev):
            text = f"* File: {item.get('file_name', 'N/A')} // {item.get('supporting_claim', 'N/A')}\n  Reason: {item.get('reasoning', '')}"
            pdf.multi_cell(0, 5, text)
            pdf.ln(1)
    else:
        pdf.cell(0, 5, "No direct strong evidence cataloged.", ln=1)
    pdf.ln(3)

    pdf.set_font('Helvetica', 'B', 10)
    pdf.cell(0, 6, "Weak / Unsupported Evidence:", ln=1)
    pdf.set_font('Helvetica', '', 9.5)
    weak_ev = report.get('weak_evidence', [])
    if weak_ev:
        for item in weak_ev:
            text = f"* File: {item.get('file_name', 'N/A')}\n  Reasoning: {item.get('reasoning', '')}"
            pdf.multi_cell(0, 5, text)
            pdf.ln(1)
    else:
        pdf.cell(0, 5, "No significant weak evidence identified.", ln=1)
    pdf.ln(3)

    pdf.set_font('Helvetica', 'B', 10)
    pdf.cell(0, 6, "Missing Evidence & Proof Requirements:", ln=1)
    pdf.set_font('Helvetica', '', 9.5)
    missing_ev = report.get('missing_evidence', [])
    if missing_ev:
        for item in missing_ev:
            text = f"* Category: {item.get('category', 'N/A')} // {item.get('description', '')}\n  Impact: {item.get('impact', '')}"
            pdf.multi_cell(0, 5, text)
            pdf.ln(1)
    else:
        pdf.cell(0, 5, "No critical missing evidence items reported.", ln=1)
    pdf.ln(6)

    # Loopholes and Contradictions
    pdf.set_font('Helvetica', 'B', 12)
    pdf.set_text_color(30, 41, 59)
    pdf.cell(0, 8, "3. CONTRACTUAL LOOPHOLES & CONTRADICTIONS", ln=1)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(2)
    
    pdf.set_font('Helvetica', 'B', 10)
    pdf.cell(0, 6, "Identified Vulnerabilities & Loopholes:", ln=1)
    pdf.set_font('Helvetica', '', 9.5)
    loopholes = report.get('loopholes', [])
    if loopholes:
        for item in loopholes:
            text = f"* {item.get('title', 'Vulnerability')} ({item.get('severity', 'Medium')} severity): {item.get('description', '')}"
            pdf.multi_cell(0, 5, text)
            pdf.ln(1)
    else:
        pdf.cell(0, 5, "No major loopholes identified in the case documentation.", ln=1)
    pdf.ln(6)

    # Courtroom Strategy and Recommendations
    pdf.set_font('Helvetica', 'B', 12)
    pdf.set_text_color(30, 41, 59)
    pdf.cell(0, 8, "4. COURTROOM STRATEGY SIMULATION", ln=1)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(2)
    
    pdf.set_font('Helvetica', 'B', 10)
    pdf.cell(0, 6, "Counterargument Risks (Opponent Attack Vectors):", ln=1)
    pdf.set_font('Helvetica', '', 9.5)
    risks = report.get('courtroom_risks', [])
    if risks:
        for item in risks:
            text = f"* {item.get('title', 'Risk')} ({item.get('severity', 'Medium')} severity): {item.get('description', '')}"
            pdf.multi_cell(0, 5, text)
            pdf.ln(1)
    else:
        pdf.cell(0, 5, "No primary counterargument risks identified.", ln=1)
    pdf.ln(3)

    pdf.set_font('Helvetica', 'B', 10)
    pdf.cell(0, 6, "Strategic Next-Step Recommendations:", ln=1)
    pdf.set_font('Helvetica', '', 9.5)
    recs = report.get('strategic_recommendations', [])
    if recs:
        for item in recs:
            pdf.multi_cell(0, 5, f"* {item}")
            pdf.ln(1)
    else:
        pdf.cell(0, 5, "No specific recommendations compiled.", ln=1)
    pdf.ln(6)

    # Legal References & Disclaimer
    pdf.set_font('Helvetica', 'B', 12)
    pdf.set_text_color(30, 41, 59)
    pdf.cell(0, 8, "5. STATUTORY & CASE LAW REFERENCES", ln=1)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(2)
    
    pdf.set_font('Helvetica', '', 9.5)
    refs = report.get('legal_references', [])
    if refs:
        for item in refs:
            text = f"* Reference: {item.get('act_or_statute', 'N/A')} - Section {item.get('section', 'N/A')}\n  Relevance: {item.get('relevance', '')}"
            pdf.multi_cell(0, 5, text)
            pdf.ln(1)
    else:
        pdf.cell(0, 5, "No specific legal references compiled for this context.", ln=1)
    pdf.ln(8)

    # Overall Case Assessment Box
    pdf.set_font('Helvetica', 'B', 10)
    pdf.cell(0, 6, "Overall Case Assessment Summary:", ln=1)
    pdf.set_font('Helvetica', '', 9.5)
    pdf.multi_cell(0, 5, str(report.get('overall_case_assessment', 'No assessment available.')))
    pdf.ln(8)

    # Disclaimer Block
    pdf.set_fill_color(255, 244, 244)
    pdf.rect(10, pdf.get_y(), 190, 20, 'F')
    pdf.set_y(pdf.get_y() + 2)
    pdf.set_x(12)
    pdf.set_font('Helvetica', 'B', 8)
    pdf.set_text_color(180, 50, 50)
    pdf.cell(0, 4, "MANDATORY AI DISCLAIMER:", ln=1)
    pdf.set_x(12)
    pdf.set_font('Helvetica', 'I', 8)
    pdf.set_text_color(120, 50, 50)
    pdf.multi_cell(186, 4, "This AI-generated analysis is intended for legal research assistance and strategic support only. It should not be treated as definitive legal advice or representation.")

    # 3. Stream File out
    pdf_bytes = pdf.output(dest='S')
    
    headers = {
        'Content-Disposition': f'attachment; filename="verdictiq_report_{workspace_id}.pdf"'
    }
    
    return StreamingResponse(
        io.BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers=headers
    )

@router.get("/docx/{workspace_id}")
async def export_docx(workspace_id: str, current_user: dict = Depends(get_user_for_export)):
    # 1. Fetch metadata and report from database
    workspace = await get_collection("workspaces").find_one({"workspace_id": workspace_id})
    report = await get_collection("agent3_final_reports").find_one({"workspace_id": workspace_id})

    if not workspace or not report:
        raise HTTPException(status_code=404, detail="Workspace metadata or final report not found. Run analysis first.")

    # 2. Build Word Document
    doc = docx.Document()
    
    # Title Section
    title = doc.add_paragraph()
    r = title.add_run("VERDICTIQ LEGAL INTELLIGENCE OS // CASE REPORT")
    r.font.name = "Arial"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(128, 128, 128)
    title.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
    
    h = doc.add_heading(level=1)
    run_h = h.add_run("LEGAL INTELLIGENCE AUDIT")
    run_h.font.name = "Arial"
    run_h.font.size = Pt(22)
    run_h.font.color.rgb = RGBColor(30, 41, 59)
    h.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    r_sub = sub.add_run(f"CASE: {workspace.get('case_title', 'Unnamed Case').upper()}")
    r_sub.font.name = "Arial"
    r_sub.font.size = Pt(14)
    r_sub.font.bold = True
    r_sub.font.color.rgb = RGBColor(79, 140, 255)
    sub.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

    # Table of case details
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Shading Accent 1'
    
    cell_data = [
        ("Case Type:", str(workspace.get('case_type', 'N/A'))),
        ("Counsel Side:", str(workspace.get('lawyer_side', 'N/A'))),
        ("Client:", str(workspace.get('client_name', 'N/A'))),
        ("Opposing Party:", str(workspace.get('opposing_party', 'N/A'))),
    ]
    
    for idx, (label, val) in enumerate(cell_data):
        row = table.rows[idx]
        row.cells[0].paragraphs[0].text = label
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[1].paragraphs[0].text = val
        
    doc.add_paragraph() # Spacing

    # 1. Executive Summary
    h1 = doc.add_heading("1. Executive Summary", level=2)
    h1.runs[0].font.color.rgb = RGBColor(30, 41, 59)
    doc.add_paragraph(str(report.get('executive_summary', 'No summary available.')))

    # 2. Evidentiary Audit
    h2 = doc.add_heading("2. Evidentiary Audit", level=2)
    h2.runs[0].font.color.rgb = RGBColor(30, 41, 59)

    doc.add_paragraph("Strong Evidence:").runs[0].font.bold = True
    strong_ev = report.get('strongest_evidence', [])
    if strong_ev:
        for item in strong_ev:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"File: {item.get('file_name', 'N/A')} // Claim: {item.get('supporting_claim', 'N/A')}\nReasoning: {item.get('reasoning', '')}")
    else:
        doc.add_paragraph("No direct strong evidence cataloged.")

    doc.add_paragraph("Weak / Unsupported Evidence:").runs[0].font.bold = True
    weak_ev = report.get('weak_evidence', [])
    if weak_ev:
        for item in weak_ev:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"File: {item.get('file_name', 'N/A')}\nReasoning: {item.get('reasoning', '')}")
    else:
        doc.add_paragraph("No significant weak evidence identified.")

    doc.add_paragraph("Missing Evidence:").runs[0].font.bold = True
    missing_ev = report.get('missing_evidence', [])
    if missing_ev:
        for item in missing_ev:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"Category: {item.get('category', 'N/A')} // {item.get('description', '')}\nImpact: {item.get('impact', '')}")
    else:
        doc.add_paragraph("No critical missing evidence reported.")

    # 3. Loopholes and Contradictions
    h3 = doc.add_heading("3. Contractual Loopholes & Contradictions", level=2)
    h3.runs[0].font.color.rgb = RGBColor(30, 41, 59)
    loopholes = report.get('loopholes', [])
    if loopholes:
        for item in loopholes:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"{item.get('title', 'Vulnerability')} ({item.get('severity', 'Medium')} severity): ").bold = True
            p.add_run(item.get('description', ''))
    else:
        doc.add_paragraph("No major loopholes identified in the case documentation.")

    # 4. Courtroom Strategy
    h4 = doc.add_heading("4. Courtroom Strategy Simulation", level=2)
    h4.runs[0].font.color.rgb = RGBColor(30, 41, 59)
    
    doc.add_paragraph("Counterargument Risks:").runs[0].font.bold = True
    risks = report.get('courtroom_risks', [])
    if risks:
        for item in risks:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"{item.get('title', 'Risk')} ({item.get('severity', 'Medium')} severity): ").bold = True
            p.add_run(item.get('description', ''))
    else:
        doc.add_paragraph("No primary counterargument risks identified.")

    doc.add_paragraph("Strategic Next-Step Recommendations:").runs[0].font.bold = True
    recs = report.get('strategic_recommendations', [])
    if recs:
        for item in recs:
            doc.add_paragraph(item, style='List Bullet')
    else:
        doc.add_paragraph("No specific recommendations compiled.")

    # 5. Statutory References
    h5 = doc.add_heading("5. Statutory & Case Law References", level=2)
    h5.runs[0].font.color.rgb = RGBColor(30, 41, 59)
    refs = report.get('legal_references', [])
    if refs:
        for item in refs:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"{item.get('act_or_statute', 'N/A')} - Section {item.get('section', 'N/A')}\n").bold = True
            p.add_run(f"Relevance: {item.get('relevance', '')}")
    else:
        doc.add_paragraph("No specific legal references compiled.")

    # Overall Assessment
    doc.add_heading("Overall Case Assessment", level=3)
    doc.add_paragraph(str(report.get('overall_case_assessment', 'No assessment available.')))

    # Disclaimer
    doc.add_paragraph()
    disc = doc.add_paragraph()
    disc.paragraph_format.left_indent = Inches(0.5)
    r_disc_title = disc.add_run("MANDATORY AI DISCLAIMER:\n")
    r_disc_title.font.bold = True
    r_disc_title.font.color.rgb = RGBColor(180, 50, 50)
    r_disc_text = disc.add_run("This AI-generated analysis is intended for legal research assistance and strategic support only. It should not be treated as definitive legal advice or representation.")
    r_disc_text.font.italic = True
    r_disc_text.font.color.rgb = RGBColor(120, 120, 120)

    # Save to memory stream
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    headers = {
        'Content-Disposition': f'attachment; filename="verdictiq_report_{workspace_id}.docx"'
    }

    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers
    )
